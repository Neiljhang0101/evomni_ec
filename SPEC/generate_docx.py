"""Convert 廣告投放工具 UX Spec markdown to DOCX for Claude Design upload."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table_from_rows(doc, rows, header_color="2B5EA7"):
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= col_count:
                break
            cell = table.cell(i, j)
            cell.text = cell_text.strip()
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
            if i == 0:
                set_cell_bg(cell, header_color)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                if cell.paragraphs and cell.paragraphs[0].runs:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                else:
                    # re-set text with bold
                    cell.paragraphs[0].clear()
                    run = cell.paragraphs[0].add_run(cell_text.strip())
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return table


def parse_table(lines, start):
    """Parse a markdown table starting at lines[start]. Returns (rows, end_index)."""
    rows = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("|"):
            if re.match(r"^\|[-\s|:]+\|$", line):
                i += 1
                continue
            cells = [c.strip() for c in line.strip("|").split("|")]
            rows.append(cells)
            i += 1
        else:
            break
    return rows, i


def process_md_to_docx(md_path, out_path):
    with open(md_path, encoding="utf-8") as f:
        content = f.read()

    lines = content.split("\n")
    doc = Document()

    # --- Global styles ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Noto Sans TC"
    font.size = Pt(10.5)

    for level, name in [(1, "Heading 1"), (2, "Heading 2"), (3, "Heading 3"), (4, "Heading 4")]:
        h = doc.styles[name]
        h.font.name = "Noto Sans TC"
        h.font.bold = True
        h.font.size = [None, Pt(20), Pt(16), Pt(13), Pt(11.5)][level]
        h.font.color.rgb = [None, RGBColor(0x1A, 0x3A, 0x5C), RGBColor(0x2B, 0x5E, 0xA7),
                             RGBColor(0x1A, 0x3A, 0x5C), RGBColor(0x40, 0x63, 0x80)][level]
        h.paragraph_format.space_before = [None, Pt(18), Pt(14), Pt(10), Pt(8)][level]
        h.paragraph_format.space_after = [None, Pt(8), Pt(6), Pt(4), Pt(3)][level]

    # --- Cover info ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("廣告投放工具（產品廣告小幫手）")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Noto Sans TC"
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
    doc.add_paragraph("Design Brief + UX Spec").alignment  # noqa

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Version: 1.0  |  Created: 2026-05-15  |  用途：上傳至 Claude Design，配合 Prompt Cheatsheet 使用").font.size = Pt(9.5)

    doc.add_paragraph()  # spacer

    # --- Parse lines ---
    i = 0
    in_code_block = False
    code_lines = []
    code_lang = ""
    current_table_rows = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code block
        if stripped.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lang = stripped[3:].strip()
                code_lines = []
                i += 1
                continue
            else:
                # End of code block — add as styled paragraph
                in_code_block = False
                code_text = "\n".join(code_lines)
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Cm(0.6)
                para.paragraph_format.space_before = Pt(4)
                para.paragraph_format.space_after = Pt(4)
                run = para.add_run(code_text)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
                # Light grey background via shading workaround
                pPr = para._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F0F2F5")
                pPr.append(shd)
                code_lines = []
                i += 1
                continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table
        if stripped.startswith("|"):
            table_rows, new_i = parse_table(lines, i)
            if table_rows:
                add_table_from_rows(doc, table_rows)
                doc.add_paragraph()  # spacer after table
                i = new_i
                continue

        # Headings
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # Remove markdown formatting from heading
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        # Horizontal rule
        if stripped in ("---", "***", "___"):
            para = doc.add_paragraph()
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "CCCCCC")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Blockquote (> ...)
        if stripped.startswith(">"):
            bq_text = stripped.lstrip(">").strip()
            # Remove emoji-like icons
            bq_text = re.sub(r"[📌✏️⚠️📊]", "", bq_text).strip()
            bq_text = re.sub(r"\*\*(.+?)\*\*", r"\1", bq_text)
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(0.8)
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(3)
            run = para.add_run(bq_text)
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            run.font.italic = True
            run.font.size = Pt(10)
            i += 1
            continue

        # List items
        list_match = re.match(r"^(\s*)[-*]\s+(.+)$", line)
        if list_match:
            indent = len(list_match.group(1))
            item_text = list_match.group(2).strip()
            item_text = re.sub(r"\*\*(.+?)\*\*", r"\1", item_text)
            item_text = re.sub(r"`(.+?)`", r"\1", item_text)
            item_text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", item_text)
            # Remove emoji
            item_text = re.sub(r"[⚠️📌✏️📊🔵🟢]", "", item_text).strip()
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent = Cm(0.6 + indent * 0.3)
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(1)
            para.add_run(item_text)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r"^(\s*)\d+\.\s+(.+)$", line)
        if num_match:
            indent = len(num_match.group(1))
            item_text = num_match.group(2).strip()
            item_text = re.sub(r"\*\*(.+?)\*\*", r"\1", item_text)
            item_text = re.sub(r"`(.+?)`", r"\1", item_text)
            para = doc.add_paragraph(style="List Number")
            para.paragraph_format.left_indent = Cm(0.6 + indent * 0.3)
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(1)
            para.add_run(item_text)
            i += 1
            continue

        # Empty line
        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        # Regular paragraph — process inline formatting
        text = stripped
        # Remove emoji
        text = re.sub(r"[📌✏️⚠️📊🔵🟢✅❌]", "", text).strip()
        if not text:
            i += 1
            continue

        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)

        # Parse bold/code/links inline
        pattern = re.compile(r"\*\*(.+?)\*\*|`(.+?)`|\[(.+?)\]\((.+?)\)")
        last = 0
        for m in pattern.finditer(text):
            if m.start() > last:
                para.add_run(text[last:m.start()])
            if m.group(1):  # bold
                run = para.add_run(m.group(1))
                run.bold = True
            elif m.group(2):  # code
                run = para.add_run(m.group(2))
                run.font.name = "Courier New"
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            elif m.group(3):  # link
                run = para.add_run(m.group(3))
                run.font.color.rgb = RGBColor(0x40, 0x9E, 0xFF)
                run.underline = True
            last = m.end()
        if last < len(text):
            para.add_run(text[last:])

        i += 1

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    import os
    base = os.path.dirname(os.path.abspath(__file__))
    md = os.path.join(base, "廣告投放工具-UX-Spec.md")
    out = os.path.join(base, "廣告投放工具-UX-Spec.docx")
    process_md_to_docx(md, out)
