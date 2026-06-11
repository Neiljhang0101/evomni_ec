"""Convert 會員中心 UX Spec markdown to DOCX for Claude Design upload."""
import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table_from_rows(doc, rows, header_color="2B5EA7"):
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= col_count:
                break
            cell = table.cell(i, j)
            cell.text = cell_text.strip()
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
            if i == 0:
                set_cell_bg(cell, header_color)
                cell.paragraphs[0].clear()
                run = cell.paragraphs[0].add_run(cell_text.strip())
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return table


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        if re.match(r"^\|[-| :]+\|$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def md_to_docx(md_path, docx_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Noto Sans TC"
    style.font.size = Pt(10)

    i = 0
    while i < len(lines):
        raw   = lines[i]
        line  = raw.rstrip("\n")
        strip = line.strip()

        # Blockquotes
        if strip.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(strip[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(0x60, 0x62, 0x66)
            run.font.size = Pt(9)
            i += 1
            continue

        # Horizontal rule
        if strip in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,4})\s+(.*)", strip)
        if m:
            level = len(m.group(1))
            text  = m.group(2)
            h_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
            p = doc.add_heading(text, level=min(level, 4))
            p.style = doc.styles[h_map.get(level, "Heading 4")]
            run = p.runs[0] if p.runs else p.add_run(text)
            sizes = {1: 18, 2: 14, 3: 12, 4: 11}
            run.font.size = Pt(sizes.get(level, 10))
            colors = {
                1: RGBColor(0x1A, 0x1B, 0x1D),
                2: RGBColor(0x2B, 0x5E, 0xA7),
                3: RGBColor(0x30, 0x31, 0x33),
                4: RGBColor(0x60, 0x62, 0x66),
            }
            run.font.color.rgb = colors.get(level, RGBColor(0x30, 0x31, 0x33))
            i += 1
            continue

        # Code blocks
        if strip.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(0.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            p.style = doc.styles["Normal"]
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x30, 0x31, 0x33)
            continue

        # Tables
        if strip.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table_from_rows(doc, rows)
            doc.add_paragraph()
            continue

        # Bullet lists
        m_ul = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if m_ul:
            indent = len(m_ul.group(1)) // 2
            text   = m_ul.group(2)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent  = Cm(0.5 + indent * 0.4)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            _add_inline(p, text)
            i += 1
            continue

        # Numbered lists
        m_ol = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if m_ol:
            indent = len(m_ol.group(1)) // 2
            text   = m_ol.group(2)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent  = Cm(0.5 + indent * 0.4)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            _add_inline(p, text)
            i += 1
            continue

        # Empty line
        if not strip:
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(4)
        _add_inline(p, strip)
        i += 1

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


def _add_inline(p, text):
    """Handle **bold**, `code`, and plain text inline."""
    parts = re.split(r"(\*\*.*?\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x2B, 0x5E, 0xA7)
        else:
            p.add_run(part)


if __name__ == "__main__":
    base = os.path.dirname(os.path.abspath(__file__))
    md   = os.path.join(base, "會員中心-UX-Spec.md")
    out  = os.path.join(base, "會員中心-UX-Spec.docx")
    md_to_docx(md, out)
